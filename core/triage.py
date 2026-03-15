"""
DONE · core/triage.py
──────────────────────
Three-layer triage system. Runs before full ingestion.

Layer 1 — Hard exclusions    (instant, no model call)
Layer 2 — Auto classification (fast scan, first page only, ~2s)
Layer 3 — Human review queue  (ambiguous files only)

Public API:
  scan_folder(folder, conn)        → queue files for triage
  scan_file(filepath, conn)        → triage a single file
  process_triage_queue(conn)       → run auto-classification
  approve(triage_id, conn)         → human approves a file
  reject(triage_id, conn, notes)   → human rejects a file
  get_pending(conn)                → files awaiting human review
  get_stats(conn)                  → triage summary stats

Usage:
  python core/triage.py --folder /path/to/folder
  python core/triage.py --process
  python core/triage.py --pending
  python core/triage.py --approve 42
  python core/triage.py --reject 42 "not relevant"
"""

import re
import sys
import time
import warnings
from pathlib import Path

warnings.filterwarnings("ignore")
warnings.filterwarnings("ignore", message=".*FloatObject.*")
warnings.filterwarnings("ignore", message=".*could not convert.*")

BASE_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(BASE_DIR / "core"))

try:
    import ollama
except ImportError:
    print("[triage] ERROR: pip install ollama")
    sys.exit(1)

from schema    import init_db, get_conn
from normalize import detect_sensitivity, detect_category, detect_quality

# ── CONFIG ────────────────────────────────────────────────────────

REASON_MODEL = "qwen2.5:14b"

# Auto-approve if quality >= this AND category is known
AUTO_APPROVE_QUALITY    = 0.65
# Auto-reject if quality <= this
AUTO_REJECT_QUALITY     = 0.15
# Min chars to be worth triaging at all
MIN_CONTENT_CHARS       = 100

SUPPORTED = {
    ".pdf", ".docx", ".doc", ".txt", ".md",
    ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
    ".csv", ".xlsx", ".xls",
    ".json", ".html"
}

# ── LAYER 1: HARD EXCLUSIONS ──────────────────────────────────────

# Directory names that are always excluded
EXCLUDED_DIRS = {
    "venv", ".venv", "env", ".env",
    "__pycache__", ".git", ".svn", ".hg",
    "node_modules", ".npm", ".yarn",
    "dist", "build", "target",
    ".pytest_cache", ".mypy_cache",
    ".tox", ".eggs", "htmlcov",
    "site-packages", "lib", "include",
    ".DS_Store", "Thumbs.db",
}

# Filename patterns that are always excluded
EXCLUDED_FILENAME_PATTERNS = [
    r'^top_level\.txt$',
    r'^entry_points\.txt$',
    r'^METADATA$',
    r'^RECORD$',
    r'^WHEEL$',
    r'^INSTALLER$',
    r'^direct_url\.json$',
    r'^\.DS_Store$',
    r'^Thumbs\.db$',
    r'^__init__\.py$',
    r'^setup\.py$',
    r'^setup\.cfg$',
    r'^pyproject\.toml$',
    r'\.egg-info$',
    r'\.pyc$',
    r'\.pyo$',
    r'\.class$',
    r'\.log$',
    r'\.tmp$',
    r'\.cache$',
    r'\.lock$',
]

# Content patterns that indicate package/system files
EXCLUDED_CONTENT_PATTERNS = [
    r'Metadata-Version:',
    r'Requires-Python:',
    r'Classifier:',
    r'^Name: \w+\nVersion:',
    r'entry_points\.txt',
    r'#!/usr/bin/',
    r'#!/usr/local/bin/',
]


def is_hard_excluded(filepath: Path) -> tuple:
    """
    Layer 1: Check if file should be excluded without any processing.
    Returns (excluded: bool, reason: str)
    """
    # Check each part of the path for excluded directories
    for part in filepath.parts:
        if part.lower() in EXCLUDED_DIRS:
            return True, f"excluded directory: {part}"

    # Check filename patterns
    name = filepath.name
    for pattern in EXCLUDED_FILENAME_PATTERNS:
        if re.search(pattern, name, re.IGNORECASE):
            return True, f"excluded filename pattern: {pattern}"

    # Check file extension
    if filepath.suffix.lower() not in SUPPORTED:
        return True, f"unsupported extension: {filepath.suffix}"

    # Check file size — skip empty files
    try:
        if filepath.stat().st_size < 10:
            return True, "file too small"
    except Exception:
        return True, "could not stat file"

    return False, ""


def check_content_excluded(text: str) -> tuple:
    """
    Check if extracted text indicates a system/package file.
    Returns (excluded: bool, reason: str)
    """
    if len(text.strip()) < 30:
        return True, f"content too short ({len(text.strip())} chars — truly empty"

    for pattern in EXCLUDED_CONTENT_PATTERNS:
        if re.search(pattern, text[:500]):
            return True, f"package/system file content detected"

    return False, ""


# ── FAST PAGE EXTRACTION ──────────────────────────────────────────

VISION_MODEL = "qwen2.5vl:7b"

def _extract_pdf_with_fallback(filepath: Path) -> str:
    """
    Extract text from a PDF.
    1. Try pypdf on all pages (fast)
    2. If total text under 30 chars, fall back to vision on pages 1-3
    3. Return best result
    """
    # Step 1: fast pypdf — try all pages, accumulate text
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(filepath))
        all_text = ""
        for page in reader.pages[:5]:
            all_text += (page.extract_text() or "")
        if len(all_text.strip()) >= 30:
            return all_text
    except Exception:
        pass

    # Step 2: scanned PDF — return placeholder so triage can still classify
    # Vision runs later during full ingestion (background, expected to be slow)
    stem = filepath.stem.replace("_", " ").replace("-", " ")
    print(f"  [triage] scanned PDF — queuing for human review: {filepath.name}")
    return f"[SCANNED PDF] {stem}"


def extract_first_page(filepath: Path) -> str:
    """
    Extract text from the first page only — fast, no model calls.
    Used for triage scanning.
    """
    ext = filepath.suffix.lower()

    try:
        if ext == ".pdf":
            text = _extract_pdf_with_fallback(filepath)
            return text[:3000]

        elif ext in {".docx", ".doc"}:
            import docx
            doc   = docx.Document(str(filepath))
            paras = [p.text for p in doc.paragraphs[:20] if p.text.strip()]
            return "\n".join(paras)[:3000]

        elif ext in {".txt", ".md"}:
            text = filepath.read_text(errors="replace")
            return text[:3000]

        elif ext in {".csv"}:
            import csv
            rows = []
            with open(filepath, newline="", errors="replace") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i >= 10:
                        break
                    rows.append(" | ".join(row))
            return "\n".join(rows)

        elif ext in {".xlsx", ".xls"}:
            try:
                import openpyxl
                wb   = openpyxl.load_workbook(str(filepath), read_only=True)
                ws   = wb.active
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 10:
                        break
                    rows.append(" | ".join(str(c) for c in row if c))
                return "\n".join(rows)
            except Exception:
                return ""

        elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
            # Return placeholder — vision model too slow for triage
            return f"[IMAGE FILE: {filepath.name}]"

        elif ext == ".json":
            import json as _json
            raw = filepath.read_text(errors="replace")
            try:
                data = _json.loads(raw)
                # Flatten to readable text
                def _flatten(obj, depth=0):
                    lines = []
                    if isinstance(obj, dict):
                        for k, v in list(obj.items())[:20]:
                            if isinstance(v, (str, int, float, bool)):
                                lines.append(f"{k}: {v}")
                            elif isinstance(v, (dict, list)):
                                lines.append(f"{k}:")
                                lines.extend(_flatten(v, depth+1))
                    elif isinstance(obj, list):
                        for item in obj[:10]:
                            lines.extend(_flatten(item, depth+1))
                    return lines
                lines = _flatten(data)
                return "\n".join(lines)[:3000]
            except Exception:
                return raw[:3000]

        elif ext == ".html":
            raw = filepath.read_text(errors="replace")
            # Strip HTML tags
            import re as _re
            text = _re.sub(r'<script[^>]*>.*?</script>', '', raw,
                           flags=_re.DOTALL | _re.IGNORECASE)
            text = _re.sub(r'<style[^>]*>.*?</style>', '', text,
                           flags=_re.DOTALL | _re.IGNORECASE)
            text = _re.sub(r'<[^>]+>', ' ', text)
            text = _re.sub(r'&nbsp;', ' ', text)
            text = _re.sub(r'&amp;', '&', text)
            text = _re.sub(r'&lt;', '<', text)
            text = _re.sub(r'&gt;', '>', text)
            text = _re.sub(r'\s+', ' ', text).strip()
            return text[:3000]

    except Exception:
        pass

    return ""


# ── DESCRIPTION GENERATION ────────────────────────────────────────

DESCRIBE_PROMPT = """Read this document text and return a JSON object with a brief description and classification.

Return ONLY valid JSON, no markdown, no backticks. Example format:
{{"description": "...", "category": "LEGAL", "sensitivity": "PUBLIC"}}

Valid categories: LEGAL, FINANCIAL, MEDICAL, INSURANCE, COMPLIANCE, ACADEMIC, TECHNICAL, PERSONAL, HR, REAL_ESTATE, GOVERNMENT, REFERENCE, JUNK, UNKNOWN
Valid sensitivity: PUBLIC, CONFIDENTIAL, PII, PRIVILEGED, HIPAA

Document filename: {filename}
Document text (first page):
{text}"""


def generate_description(filepath: Path, text: str) -> dict:
    """
    Generate a 1-2 sentence description and classification
    using qwen2.5:14b. Called only for non-obvious files.
    """
    prompt = DESCRIBE_PROMPT.format(
        filename=filepath.name,
        text=text[:2000]
    )

    try:
        response = ollama.chat(
            model=REASON_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.1, "num_ctx": 2048}
        )
        raw = response["message"]["content"].strip()
        raw = re.sub(r'```(?:json)?\s*', '', raw).strip()

        start = raw.find('{')
        end   = raw.rfind('}')
        if start == -1 or end == -1:
            return {}

        return __import__('json').loads(raw[start:end + 1])

    except Exception:
        return {}


# ── DUPLICATE DETECTION ───────────────────────────────────────────

def check_duplicate(doc_hash: str, conn) -> int:
    """
    Check if this file hash already exists in documents.
    Returns doc_id if duplicate, 0 if not.
    """
    row = conn.execute(
        "SELECT id FROM documents WHERE doc_hash = ?",
        (doc_hash,)
    ).fetchone()
    return row["id"] if row else 0


# ── PAGE COUNT ────────────────────────────────────────────────────

def get_page_count(filepath: Path) -> int:
    """Fast page count without full extraction."""
    ext = filepath.suffix.lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            return len(PdfReader(str(filepath)).pages)
        elif ext in {".docx", ".doc"}:
            import docx
            doc = docx.Document(str(filepath))
            return max(1, len(doc.paragraphs) // 40)
        elif ext in {".txt", ".md"}:
            chars = len(filepath.read_text(errors="replace"))
            return max(1, chars // 3000)
        else:
            return 1
    except Exception:
        size_mb = filepath.stat().st_size / (1024 * 1024)
        return max(1, int(size_mb * 20))


def estimate_seconds(page_count: int) -> int:
    """Estimate full ingestion time in seconds."""
    batches = max(1, (page_count + 2) // 3)
    return batches * 15


def get_priority(page_count: int) -> int:
    if page_count <= 10:   return 1
    if page_count <= 50:   return 2
    if page_count <= 200:  return 3
    return 4


# ── CORE TRIAGE ───────────────────────────────────────────────────

def scan_file(filepath, conn) -> dict:
    """
    Triage a single file through all three layers.
    Returns result dict with decision and reason.
    """
    import hashlib
    filepath = Path(filepath).resolve()
    now      = time.strftime("%Y-%m-%dT%H:%M:%SZ")

    # Already in triage queue?
    existing = conn.execute(
        "SELECT id, decision FROM triage_queue WHERE filepath = ?",
        (str(filepath),)
    ).fetchone()
    if existing:
        return {
            "action":   "skipped",
            "reason":   f"already in triage (decision={existing['decision']})",
            "filepath": str(filepath),
        }

    # Already fully indexed?
    try:
        with open(filepath, "rb") as f:
            doc_hash = hashlib.sha256(f.read()).hexdigest()
    except Exception:
        return {"action": "error", "reason": "could not read file",
                "filepath": str(filepath)}

    dup_id = check_duplicate(doc_hash, conn)

    # ── Layer 1: Hard exclusion ────────────────────────────────
    excluded, reason = is_hard_excluded(filepath)
    if excluded:
        return {
            "action":   "excluded",
            "reason":   reason,
            "filepath": str(filepath),
        }

    # ── Extract first page ─────────────────────────────────────
    text       = extract_first_page(filepath)
    page_count = get_page_count(filepath)
    file_size  = filepath.stat().st_size
    priority   = get_priority(page_count)
    est_secs   = estimate_seconds(page_count)

    # ── Content exclusion ──────────────────────────────────────
    content_excluded, content_reason = check_content_excluded(text)
    if content_excluded and not dup_id:
        # Store as auto-rejected in triage queue
        conn.execute("""
            INSERT OR IGNORE INTO triage_queue
            (filepath, filename, file_type, file_size, page_count,
             estimated_secs, priority, description, category,
             sensitivity, quality_score, sample_text,
             is_duplicate, duplicate_of,
             decision, decided_by, decided_at, created_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            str(filepath), filepath.name, filepath.suffix.lower(),
            file_size, page_count, est_secs, priority,
            content_reason, "JUNK", "PUBLIC", 0.0,
            text[:300], 0, 0,
            "rejected", "auto", now, now
        ))
        conn.commit()
        return {
            "action":   "auto_rejected",
            "reason":   content_reason,
            "filepath": str(filepath),
        }

    # ── Layer 2: Auto classification ──────────────────────────
    is_scanned  = text.startswith("[SCANNED PDF]")
    quality     = detect_quality(text, page_count) if not is_scanned else 0.5
    category    = detect_category(text, filepath.name)
    sensitivity = detect_sensitivity(text)

    # Determine decision
    if dup_id:
        decision   = "rejected"
        decided_by = "auto"
        desc       = f"Duplicate of document ID {dup_id}"
        is_dup     = 1
    elif is_scanned:
        # Scanned PDFs always go to human review — vision runs at ingest time
        decision   = "pending"
        decided_by = "pending"
        desc       = ""
        is_dup     = 0
    elif quality <= AUTO_REJECT_QUALITY:
        decision   = "rejected"
        decided_by = "auto"
        desc       = f"Low quality content (score={quality:.2f})"
        is_dup     = 0
    elif quality >= AUTO_APPROVE_QUALITY and category != "UNKNOWN":
        decision   = "approved"
        decided_by = "auto"
        desc       = ""  # will generate below
        is_dup     = 0
    else:
        decision   = "pending"
        decided_by = "pending"
        desc       = ""
        is_dup     = 0

    # Generate description for approved/pending files
    if decision in ("approved", "pending") and text:
        result = generate_description(filepath, text)
        if result:
            desc        = result.get("description", desc)
            category    = result.get("category", category)
            sensitivity = result.get("sensitivity", sensitivity)

    # Store in triage queue
    conn.execute("""
        INSERT OR IGNORE INTO triage_queue
        (filepath, filename, file_type, file_size, page_count,
         estimated_secs, priority, description, category,
         sensitivity, quality_score, sample_text,
         is_duplicate, duplicate_of,
         decision, decided_by, decided_at, created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        str(filepath), filepath.name, filepath.suffix.lower(),
        file_size, page_count, est_secs, priority,
        desc, category, sensitivity, quality,
        text[:500], is_dup, dup_id,
        decision, decided_by,
        now if decision != "pending" else "", now
    ))
    conn.commit()

    # If auto-approved, queue for ingestion
    if decision == "approved":
        _queue_for_ingest(filepath, page_count, file_size,
                          priority, est_secs, conn)

    return {
        "action":      decision,
        "reason":      desc,
        "filepath":    str(filepath),
        "category":    category,
        "sensitivity": sensitivity,
        "quality":     quality,
        "pages":       page_count,
        "priority":    priority,
    }


def _queue_for_ingest(filepath, page_count, file_size,
                      priority, est_secs, conn):
    """Add an approved file to the ingest job queue."""
    now = time.strftime("%Y-%m-%dT%H:%M:%SZ")

    # Don't double-queue
    existing = conn.execute(
        "SELECT id FROM ingest_jobs WHERE filepath = ? AND status != 'cancelled'",
        (str(filepath),)
    ).fetchone()
    if existing:
        return

    conn.execute("""
        INSERT INTO ingest_jobs
        (filepath, filename, status, priority, page_count,
         file_size_mb, estimated_secs, created_at)
        VALUES (?,?,?,?,?,?,?,?)
    """, (
        str(filepath), Path(filepath).name,
        "pending", priority, page_count,
        round(file_size / (1024 * 1024), 3),
        est_secs, now
    ))
    conn.commit()


def scan_folder(folder, conn, verbose=True) -> dict:
    """
    Scan all files in a folder through the triage pipeline.
    Returns summary of decisions made.
    """
    folder = Path(folder)
    if not folder.exists():
        return {"error": f"Folder not found: {folder}"}

    # Collect all supported files
    files = []
    for ext in SUPPORTED:
        files.extend(folder.rglob(f"*{ext}"))

    if not files:
        return {"error": f"No supported files found in {folder}"}

    if verbose:
        print(f"[triage] Found {len(files)} files in {folder}")
        print(f"[triage] Running triage...\n")

    counts = {
        "excluded":      0,
        "auto_rejected": 0,
        "auto_approved": 0,
        "pending":       0,
        "skipped":       0,
        "error":         0,
    }

    for filepath in files:
        result = scan_file(filepath, conn)
        action = result.get("action", "error")

        if action == "approved":
            counts["auto_approved"] += 1
        elif action in counts:
            counts[action] += 1
        else:
            counts["error"] += 1

        if verbose:
            icon = {
                "excluded":      "—",
                "auto_rejected": "✗",
                "approved":      "✓",
                "pending":       "?",
                "skipped":       "·",
                "error":         "!",
            }.get(action, "?")
            print(f"  {icon} {Path(filepath).name[:50]:<50} "
                  f"{action:<15} {result.get('reason', '')[:40]}")

    total = len(files)
    if verbose:
        print(f"\n[triage] ── Summary ──")
        print(f"  Total files:    {total}")
        print(f"  Auto-approved:  {counts['auto_approved']}  → queued for ingestion")
        print(f"  Pending review: {counts['pending']}  → needs your decision")
        print(f"  Auto-rejected:  {counts['auto_rejected']}  → junk/low quality")
        print(f"  Hard excluded:  {counts['excluded']}  → system files")
        print(f"  Skipped:        {counts['skipped']}  → already processed")

    counts["total"] = total
    return counts


# ── HUMAN REVIEW ──────────────────────────────────────────────────

def get_pending(conn, owner_id=1, limit=50) -> list:
    """Return files awaiting human review."""
    rows = conn.execute("""
        SELECT id, filename, file_type, file_size, page_count,
               estimated_secs, priority, description, category,
               sensitivity, quality_score, sample_text,
               is_duplicate, duplicate_of, created_at
        FROM triage_queue
        WHERE decision = 'pending'
          AND owner_id = ?
        ORDER BY priority ASC, quality_score DESC
        LIMIT ?
    """, (owner_id, limit)).fetchall()
    return [dict(r) for r in rows]


def approve(triage_id: int, conn, notes: str = "") -> bool:
    """Human approves a file — queues it for ingestion."""
    now = time.strftime("%Y-%m-%dT%H:%M:%SZ")

    row = conn.execute(
        "SELECT * FROM triage_queue WHERE id = ?", (triage_id,)
    ).fetchone()

    if not row:
        return False

    conn.execute("""
        UPDATE triage_queue
        SET decision='approved', decided_by='human',
            decided_at=?, decision_notes=?
        WHERE id=?
    """, (now, notes, triage_id))
    conn.commit()

    # Queue for ingestion
    _queue_for_ingest(
        row["filepath"], row["page_count"], row["file_size"],
        row["priority"], row["estimated_secs"], conn
    )
    return True


def reject(triage_id: int, conn, notes: str = "") -> bool:
    """Human rejects a file — marks as rejected, not indexed."""
    now = time.strftime("%Y-%m-%dT%H:%M:%SZ")

    result = conn.execute("""
        UPDATE triage_queue
        SET decision='rejected', decided_by='human',
            decided_at=?, decision_notes=?
        WHERE id=?
    """, (now, notes, triage_id))
    conn.commit()
    return result.rowcount > 0


def get_stats(conn, owner_id=1) -> dict:
    """Return triage summary statistics."""
    rows = conn.execute("""
        SELECT decision, decided_by, COUNT(*) as c
        FROM triage_queue
        WHERE owner_id = ?
        GROUP BY decision, decided_by
    """, (owner_id,)).fetchall()

    stats = {
        "total":          0,
        "pending":        0,
        "approved":       0,
        "rejected":       0,
        "auto_approved":  0,
        "auto_rejected":  0,
        "human_approved": 0,
        "human_rejected": 0,
    }

    for r in rows:
        decision   = r["decision"]
        decided_by = r["decided_by"]
        count      = r["c"]

        stats["total"] += count

        if decision == "pending":
            stats["pending"] += count
        elif decision == "approved":
            stats["approved"] += count
            if decided_by == "auto":
                stats["auto_approved"] += count
            else:
                stats["human_approved"] += count
        elif decision == "rejected":
            stats["rejected"] += count
            if decided_by == "auto":
                stats["auto_rejected"] += count
            else:
                stats["human_rejected"] += count

    return stats


# ── CLI ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    p = argparse.ArgumentParser(description="DONE triage — smart file classification")
    p.add_argument("--folder",  help="Scan a folder for triage")
    p.add_argument("--file",    help="Triage a single file")
    p.add_argument("--pending", action="store_true",
                   help="Show files awaiting human review")
    p.add_argument("--approve", type=int, metavar="ID",
                   help="Approve a triage item by ID")
    p.add_argument("--reject",  type=int, metavar="ID",
                   help="Reject a triage item by ID")
    p.add_argument("--notes",   default="",
                   help="Notes for approve/reject decision")
    p.add_argument("--stats",   action="store_true",
                   help="Show triage statistics")
    args = p.parse_args()

    conn = init_db()

    if args.folder:
        scan_folder(args.folder, conn, verbose=True)

    elif args.file:
        result = scan_file(args.file, conn)
        print(f"\n── Triage result: {Path(args.file).name} ──")
        for k, v in result.items():
            print(f"  {k:<12} {v}")

    elif args.pending:
        items = get_pending(conn)
        if not items:
            print("[triage] No files pending review")
        else:
            print(f"\n── {len(items)} files pending review ──\n")
            for item in items:
                mins = item['estimated_secs'] // 60
                secs = item['estimated_secs'] % 60
                print(f"  ID {item['id']:<4} [{item['category']:<12}] "
                      f"{item['filename'][:45]}")
                print(f"         {item['description'][:70]}")
                print(f"         {item['page_count']} pages  "
                      f"quality={item['quality_score']:.2f}  "
                      f"~{mins}m{secs}s  "
                      f"sensitivity={item['sensitivity']}")
                print(f"         python core/triage.py --approve {item['id']}")
                print()

    elif args.approve:
        if approve(args.approve, conn, args.notes):
            print(f"[triage] ✓ Approved ID {args.approve} — queued for ingestion")
        else:
            print(f"[triage] ID {args.approve} not found")

    elif args.reject:
        if reject(args.reject, conn, args.notes):
            print(f"[triage] ✗ Rejected ID {args.reject}")
        else:
            print(f"[triage] ID {args.reject} not found")

    elif args.stats:
        stats = get_stats(conn)
        print(f"\n── Triage Stats ──")
        for k, v in stats.items():
            print(f"  {k:<20} {v}")

    else:
        p.print_help()