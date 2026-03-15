"""
DONE · core/ingest.py  (v2)
────────────────────────────
Page-aware document ingestion pipeline.

Reads any document type, processes page by page, extracts entities
in batches of 3 pages, chunks with size tuned to document length,
embeds into ChromaDB with page metadata.

Queue-ready: every public function returns structured results so the
job queue can track progress, estimate time, and prioritise small
files over large ones.

Public API:
  estimate_job(filepath)           -> priority, page count, ETA
  ingest_file(filepath, conn)      -> full ingestion, returns result dict
  list_documents(conn)             -> all indexed documents
  init_db()                        -> create/migrate schema

Usage:
  python core/ingest.py --file path/to/doc.pdf --preview
  python core/ingest.py --folder path/to/folder
  python core/ingest.py --list
  python core/ingest.py --estimate path/to/doc.pdf
"""

import argparse
import base64
import hashlib
import json
import re
import sqlite3
import sys
import time
from pathlib import Path

try:
    import ollama
except ImportError:
    print("[ingest] ERROR: pip install ollama")
    sys.exit(1)

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import (
        VectorParams, Distance,
        PointStruct,
        Filter, FieldCondition, MatchValue, FilterSelector,
    )
except ImportError:
    print("[ingest] ERROR: pip install qdrant-client")
    sys.exit(1)

# ── CONFIG ────────────────────────────────────────────────────────
BASE_DIR     = Path(__file__).parent.parent
DB_PATH      = BASE_DIR / "db" / "done.db"
QDRANT_PATH  = BASE_DIR / "db" / "qdrant"

VISION_MODEL  = "qwen2.5vl:7b"
REASON_MODEL  = "qwen2.5:14b"   # QA answering only — keep quality high
ENTITY_MODEL  = "llama3.1:8b"   # Entity extraction — 4-5x faster than 14b, sufficient for NER
EMBED_MODEL   = "nomic-embed-text"
EMBED_URL     = "http://localhost:11434/api/embeddings"

# Pages with fewer avg chars than this are treated as slide decks —
# entity extraction is skipped (nothing useful to extract).
SPARSE_THRESHOLD = 300

# Documents longer than this skip entity extraction regardless of density.
# A 600-page textbook doesn't benefit enough to justify the runtime cost.
MAX_ENTITY_PAGES = 100

CHUNK_PROFILE = {
    "small":  {"size": 600, "overlap": 100, "max_pages": 10},
    "medium": {"size": 400, "overlap": 80,  "max_pages": 50},
    "large":  {"size": 300, "overlap": 60,  "max_pages": 99999},
}

ENTITY_BATCH_SIZE = 3
SECONDS_PER_BATCH = 4   # llama3.1:8b is ~4s/batch vs 15s for qwen2.5:14b

SUPPORTED = {
    ".pdf", ".docx", ".doc", ".txt", ".md",
    ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
    ".csv", ".xlsx", ".xls",
    ".json", ".html"
}

# ── SCHEMA ────────────────────────────────────────────────────────

SCHEMA = """
CREATE TABLE IF NOT EXISTS documents (
    id            INTEGER PRIMARY KEY AUTOINCREMENT,
    filename      TEXT NOT NULL,
    filepath      TEXT NOT NULL UNIQUE,
    file_type     TEXT,
    file_size     INTEGER,
    page_count    INTEGER DEFAULT 0,
    date_modified TEXT,
    date_indexed  TEXT,
    raw_text      TEXT,
    doc_hash      TEXT UNIQUE,
    chunk_count   INTEGER DEFAULT 0,
    entity_count  INTEGER DEFAULT 0,
    status        TEXT DEFAULT 'indexed'
);

CREATE TABLE IF NOT EXISTS entities (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    doc_id       INTEGER REFERENCES documents(id) ON DELETE CASCADE,
    entity_type  TEXT,
    value        TEXT,
    context      TEXT,
    page_number  INTEGER DEFAULT 0,
    confidence   REAL DEFAULT 1.0,
    created_at   TEXT
);

CREATE TABLE IF NOT EXISTS chunks (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    doc_id       INTEGER REFERENCES documents(id) ON DELETE CASCADE,
    chunk_index  INTEGER,
    text         TEXT,
    page_start   INTEGER DEFAULT 0,
    page_end     INTEGER DEFAULT 0,
    chroma_id    TEXT
);

CREATE TABLE IF NOT EXISTS ingest_jobs (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    filepath        TEXT NOT NULL,
    filename        TEXT NOT NULL,
    status          TEXT DEFAULT 'pending',
    priority        INTEGER DEFAULT 2,
    page_count      INTEGER DEFAULT 0,
    file_size_mb    REAL DEFAULT 0,
    estimated_secs  INTEGER DEFAULT 0,
    progress_page   INTEGER DEFAULT 0,
    error_message   TEXT,
    created_at      TEXT,
    started_at      TEXT,
    finished_at     TEXT
);
"""


def _open_db(path=None, timeout=30):
    """Open a SQLite connection with WAL mode and row factory."""
    p = path or DB_PATH
    Path(p).parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(p), timeout=timeout, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=30000")
    conn.execute("PRAGMA synchronous=NORMAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db():
    conn = _open_db()
    conn.executescript(SCHEMA)
    conn.commit()
    return conn


# ── JOB ESTIMATION ────────────────────────────────────────────────

def estimate_job(filepath):
    filepath  = Path(filepath)
    if not filepath.exists():
        return {"error": f"File not found: {filepath}"}

    ext      = filepath.suffix.lower()
    size_mb  = filepath.stat().st_size / (1024 * 1024)
    page_count = 0

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            page_count = len(PdfReader(str(filepath)).pages)
        except Exception:
            page_count = max(1, int(size_mb * 20))
    elif ext in {".docx", ".doc"}:
        try:
            import docx
            doc = docx.Document(str(filepath))
            page_count = max(1, len(doc.paragraphs) // 40)
        except Exception:
            page_count = max(1, int(size_mb * 10))
    elif ext in {".txt", ".md"}:
        try:
            chars = len(filepath.read_text(errors="replace"))
            page_count = max(1, chars // 3000)
        except Exception:
            page_count = 1
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        page_count = 1
    else:
        page_count = max(1, int(size_mb * 5))

    # Entity extraction is skipped for large docs — don't include that time
    if page_count > MAX_ENTITY_PAGES:
        n_batches      = 0
        estimated_secs = max(5, int(page_count * 0.3))  # extraction + embedding only
    else:
        n_batches      = max(1, (page_count + ENTITY_BATCH_SIZE - 1) // ENTITY_BATCH_SIZE)
        estimated_secs = max(5, int(page_count * 0.3)) + (n_batches * SECONDS_PER_BATCH)

    if page_count <= 10:
        priority = 1
    elif page_count <= 50:
        priority = 2
    elif page_count <= 200:
        priority = 3
    else:
        priority = 4

    if page_count <= CHUNK_PROFILE["small"]["max_pages"]:
        profile = "small"
    elif page_count <= CHUNK_PROFILE["medium"]["max_pages"]:
        profile = "medium"
    else:
        profile = "large"

    return {
        "filepath":       str(filepath),
        "filename":       filepath.name,
        "file_type":      ext,
        "size_mb":        round(size_mb, 2),
        "page_count":     page_count,
        "priority":       priority,
        "estimated_secs": estimated_secs,
        "chunk_profile":  profile,
        "n_batches":      n_batches,
    }


# ── PAGE EXTRACTION ───────────────────────────────────────────────

def read_pdf_pages(filepath):
    """Extract text from PDF. Uses PyMuPDF (fast) with pypdf fallback."""
    import importlib.util
    if importlib.util.find_spec("fitz") is not None:
        return _read_pdf_fitz(filepath)
    return _read_pdf_pypdf(filepath)


def _read_pdf_fitz(filepath):
    import fitz
    doc   = fitz.open(str(filepath))
    pages = []
    for i in range(len(doc)):
        page_num = i + 1
        try:
            text = doc[i].get_text("text") or ""
        except Exception as e:
            print(f"  [page {page_num}] fitz error ({type(e).__name__}), skipping")
            continue

        if len(text.strip()) < 30:
            print(f"  [page {page_num}] scanned, using vision...", end=" ", flush=True)
            text = read_page_with_vision(filepath, page_num) or ""
            print(f"{len(text)} chars")
        else:
            print(f"  [page {page_num}] {len(text)} chars")

        if text.strip():
            pages.append((page_num, text))

    doc.close()
    return pages


def _read_pdf_pypdf(filepath):
    import warnings
    from pypdf import PdfReader
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        reader = PdfReader(str(filepath))
    pages  = []

    for i in range(len(reader.pages)):
        page_num = i + 1
        try:
            text = reader.pages[i].extract_text() or ""
        except Exception as e:
            print(f"  [page {page_num}] extract error ({type(e).__name__}), skipping")
            continue

        if len(text.strip()) < 30:
            print(f"  [page {page_num}] scanned, using vision...", end=" ", flush=True)
            text = read_page_with_vision(filepath, page_num) or ""
            print(f"{len(text)} chars")
        else:
            print(f"  [page {page_num}] {len(text)} chars")

        if text.strip():
            pages.append((page_num, text))

    return pages


def read_docx_pages(filepath):
    import docx
    doc   = docx.Document(str(filepath))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]

    page_size = 40
    pages     = []
    for i in range(0, len(paras), page_size):
        page_num = (i // page_size) + 1
        text     = "\n\n".join(paras[i:i + page_size])
        pages.append((page_num, text))
        print(f"  [page {page_num}] {len(text)} chars")

    return pages


def read_txt_pages(filepath):
    text  = Path(filepath).read_text(errors="replace")
    lines = text.splitlines()
    page_size = 100
    pages = []
    for i in range(0, max(1, len(lines)), page_size):
        page_num = (i // page_size) + 1
        chunk    = "\n".join(lines[i:i + page_size])
        if chunk.strip():
            pages.append((page_num, chunk))
    return pages or [(1, text)]


def read_csv_pages(filepath):
    import csv
    rows = []
    with open(filepath, newline="", errors="replace") as f:
        for row in csv.reader(f):
            rows.append(" | ".join(row))
    return [(1, "\n".join(rows))]


def read_image_page(filepath):
    b64  = base64.b64encode(open(filepath, "rb").read()).decode()
    response = ollama.chat(
        model=VISION_MODEL,
        messages=[{
            "role": "user",
            "content": (
                "Read this document completely and accurately. "
                "Extract ALL text you can see, preserving structure. "
                "If handwritten, read it carefully. "
                "Output the full text content only."
            ),
            "images": [b64],
        }]
    )
    return [(1, response["message"]["content"])]


def read_page_with_vision(pdf_path, page_num):
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(
            str(pdf_path), first_page=page_num,
            last_page=page_num, dpi=150
        )
        if images:
            import io
            buf = io.BytesIO()
            images[0].save(buf, format="JPEG")
            b64 = base64.b64encode(buf.getvalue()).decode()
            response = ollama.chat(
                model=VISION_MODEL,
                messages=[{
                    "role": "user",
                    "content": f"Read page {page_num} of this document. Output text only.",
                    "images": [b64],
                }]
            )
            return response["message"]["content"]
    except ImportError:
        pass
    except Exception as e:
        print(f"  [vision] Error on page {page_num}: {e}")
    return ""


def extract_pages(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return read_pdf_pages(filepath)
    elif ext in {".docx", ".doc"}:
        return read_docx_pages(filepath)
    elif ext in {".txt", ".md"}:
        return read_txt_pages(filepath)
    elif ext in {".csv"}:
        return read_csv_pages(filepath)
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        return read_image_page(filepath)
    elif ext == ".json":
        import json as _json
        try:
            raw  = Path(filepath).read_text(errors="replace")
            data = _json.loads(raw)

            def _flatten_obj(obj):
                """Convert a dict or scalar to readable key: value lines."""
                lines = []
                if isinstance(obj, dict):
                    for k, v in list(obj.items())[:50]:
                        if isinstance(v, (str, int, float, bool)):
                            lines.append(f"{k}: {v}")
                        elif isinstance(v, (dict, list)):
                            sub = _flatten_obj(v)
                            if sub:
                                lines.append(f"{k}: {sub}")
                elif isinstance(obj, list):
                    lines.append(", ".join(str(x) for x in obj[:30]))
                else:
                    lines.append(str(obj))
                return " | ".join(lines)

            # Array of objects → one page per entry for clean per-record chunking
            if isinstance(data, list):
                pages = []
                for i, item in enumerate(data[:500]):
                    text = _flatten_obj(item) if isinstance(item, dict) else str(item)
                    if text.strip():
                        pages.append((i + 1, text))
                return pages or [(1, raw[:5000])]

            # Top-level dict → single page
            text = _flatten_obj(data)
            return [(1, text)] if text.strip() else [(1, raw[:5000])]
        except Exception:
            return [(1, Path(filepath).read_text(errors="replace")[:5000])]
    elif ext == ".html":
        import re as _re
        raw  = Path(filepath).read_text(errors="replace")
        text = _re.sub(r'<script[^>]*>.*?</script>', '', raw,
                       flags=_re.DOTALL | _re.IGNORECASE)
        text = _re.sub(r'<style[^>]*>.*?</style>', '', text,
                       flags=_re.DOTALL | _re.IGNORECASE)
        text = _re.sub(r'<[^>]+>', ' ', text)
        text = _re.sub(r'&nbsp;', ' ', text)
        text = _re.sub(r'&amp;', '&', text)
        text = _re.sub(r'&lt;', '<', text)
        text = _re.sub(r'&gt;', '>', text)
        text = _re.sub(r'\s+', ' ', text).strip()
        return [(1, text)]
    else:
        try:
            text = Path(filepath).read_text(errors="replace")
            return [(1, text)]
        except Exception:
            return []


# ── CHUNKING ──────────────────────────────────────────────────────

def chunk_pages(pages, profile="small"):
    cfg     = CHUNK_PROFILE[profile]
    size    = cfg["size"]
    overlap = cfg["overlap"]
    chunks  = []
    idx     = 0

    for page_num, page_text in pages:
        words = page_text.split()
        i     = 0
        while i < len(words):
            text = " ".join(words[i:i + size])
            if len(text.strip()) > 50:
                chunks.append({
                    "text":        text,
                    "page_start":  page_num,
                    "page_end":    page_num,
                    "chunk_index": idx,
                })
                idx += 1
            i += size - overlap

    return chunks


# ── ENTITY EXTRACTION ─────────────────────────────────────────────

ENTITY_PROMPT = """You are an entity extractor. Extract ALL named entities from the document text below.

Find and return entities for these types:
- PERSON: any person's name
- ORG: any company, university, or organization name
- LOCATION: any city, state, address, or place
- CONTACT: any phone number, email, or website URL
- DATE: any date, year, or time period
- AMOUNT: any dollar amount, salary, or numeric quantity
- SKILL: any technical skill, tool, or technology mentioned
- CLAUSE: any important terms, conditions, or policy statements

IMPORTANT: There will almost always be entities in a real document. Look carefully.

Return ONLY a valid JSON array. No markdown, no backticks, no explanation.
Just the array starting with [ and ending with ].

Example:
[
  {"type": "PERSON", "value": "John Smith", "context": "John Smith works at Google"},
  {"type": "ORG", "value": "Google", "context": "John Smith works at Google"}
]

If truly nothing found return: []"""


def generate_doc_questions(filename, chunks, conn, doc_id, model=None):
    """
    Generate 4 specific questions a user might ask about this document.
    Questions explicitly include the filename so the user knows which doc to reference.
    Stores results as a JSON array in documents.questions.
    """
    model = model or ENTITY_MODEL

    # Use top 3 chunks as context — enough to understand the doc without overloading the prompt
    context_chunks = chunks[:3]
    context = "\n\n".join(
        f"[Page {c.get('page_start', '?')}]\n{c.get('text', '')[:800]}"
        for c in context_chunks
    )
    if not context.strip():
        return []

    prompt = f"""You are analyzing a document called "{filename}".

Read the following content and generate exactly 4 specific questions a user would ask about this document.

Rules:
- Every question MUST contain the exact filename: {filename}
- Use natural phrasing like "In {filename}, what...", "What does {filename} say about...", "According to {filename}, who..."
- Questions should be genuinely useful and specific to the content
- No numbering, no bullet points, no dashes — just the 4 questions, one per line

Document content:
{context}

Output exactly 4 questions, one per line:"""

    try:
        resp = ollama.chat(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.4, "num_predict": 300},
        )
        raw = resp.message.content.strip()
        questions = [
            q.strip().lstrip("0123456789.-) ")
            for q in raw.splitlines()
            if q.strip() and "?" in q
        ][:4]

        if questions:
            conn.execute(
                "UPDATE documents SET questions=? WHERE id=?",
                (json.dumps(questions), doc_id)
            )
            conn.commit()
            print(f"[ingest] ✓ generated {len(questions)} questions")

        return questions

    except Exception as e:
        print(f"[ingest] Question generation skipped: {e}")
        return []


def extract_entities_from_batch(page_batch, doc_id, conn):
    combined = ""
    for page_num, text in page_batch:
        combined += f"\n[PAGE {page_num}]\n{text[:1500]}\n"

    try:
        response = ollama.chat(
            model=ENTITY_MODEL,
            messages=[{
                "role": "user",
                "content": f"{ENTITY_PROMPT}\n\nDocument text:\n{combined}"
            }],
            options={"temperature": 0.1, "num_ctx": 2048}
        )
        raw = response["message"]["content"].strip()
        raw = re.sub(r'```(?:json)?\s*', '', raw).strip()

        start = raw.find('[')
        end   = raw.rfind(']')
        if start == -1 or end == -1:
            return 0

        entities = json.loads(raw[start:end + 1])
        if not isinstance(entities, list):
            return 0

        default_page = page_batch[0][0] if page_batch else 0
        now          = time.strftime("%Y-%m-%dT%H:%M:%SZ")
        count        = 0

        for e in entities:
            if not all(k in e for k in ("type", "value", "context")):
                continue

            page_num = default_page
            for pn, pt in page_batch:
                if str(e.get("value", "")).lower() in pt.lower():
                    page_num = pn
                    break

            conn.execute(
                "INSERT INTO entities "
                "(doc_id, entity_type, value, context, page_number, created_at, updated_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (doc_id, e["type"], str(e["value"])[:500],
                 str(e["context"])[:500], page_num, now, now)
            )
            count += 1

        conn.commit()
        return count

    except json.JSONDecodeError:
        return 0
    except Exception as e:
        print(f"  [entity] Error: {e}")
        return 0


# ── EMBEDDING ─────────────────────────────────────────────────────

def _ensure_collection(client):
    """Create the Qdrant collection if it doesn't exist yet."""
    existing = {c.name for c in client.get_collections().collections}
    if "done_docs" not in existing:
        client.create_collection(
            collection_name="done_docs",
            vectors_config={
                "dense": VectorParams(size=768, distance=Distance.COSINE)
            },
        )


def embed_chunks(chunks, doc_id, filename, conn):
    if not chunks:
        return 0

    QDRANT_PATH.mkdir(parents=True, exist_ok=True)
    client = QdrantClient(path=str(QDRANT_PATH))
    _ensure_collection(client)

    # Delete stale points for this doc (handles re-ingest)
    try:
        client.delete(
            collection_name="done_docs",
            points_selector=FilterSelector(
                filter=Filter(must=[
                    FieldCondition(key="doc_id", match=MatchValue(value=int(doc_id)))
                ])
            ),
        )
    except Exception:
        pass

    texts      = [c["text"] for c in chunks]
    dense_vecs = [ollama.embeddings(model=EMBED_MODEL, prompt=t)["embedding"] for t in texts]

    now    = time.strftime("%Y-%m-%dT%H:%M:%SZ")
    points = []

    for c, dense in zip(chunks, dense_vecs):
        point_id = int(doc_id) * 100_000 + int(c["chunk_index"])
        points.append(PointStruct(
            id=point_id,
            vector={"dense": dense},
            payload={
                "doc_id":      int(doc_id),
                "filename":    filename,
                "chunk_index": int(c["chunk_index"]),
                "page_start":  int(c["page_start"]),
                "page_end":    int(c["page_end"]),
                "text":        c["text"],
            },
        ))
        conn.execute(
            "INSERT OR REPLACE INTO chunks "
            "(doc_id, chunk_index, text, page_start, page_end, chroma_id, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (doc_id, c["chunk_index"], c["text"],
             c["page_start"], c["page_end"], str(point_id), now)
        )

    conn.commit()
    client.upsert(collection_name="done_docs", points=points)
    return len(points)


# ── MAIN INGEST ───────────────────────────────────────────────────

def ingest_file(filepath, conn, preview=False, job_id=None):
    filepath = Path(filepath).resolve()

    if not filepath.exists():
        return {"success": False, "error": f"File not found: {filepath}"}

    if filepath.suffix.lower() not in SUPPORTED:
        return {"success": False, "error": f"Unsupported type: {filepath.suffix}"}

    with open(filepath, "rb") as f:
        doc_hash = hashlib.sha256(f.read()).hexdigest()

    existing = conn.execute(
        "SELECT id, filename FROM documents WHERE doc_hash = ?",
        (doc_hash,)
    ).fetchone()

    if existing:
        print(f"[ingest] Already indexed: {filepath.name} (id={existing['id']})")
        return {"success": True, "skipped": True,
                "doc_id": existing["id"], "filename": filepath.name}

    estimate = estimate_job(filepath)
    profile  = estimate["chunk_profile"]

    print(f"\n[ingest] ── {filepath.name} ──")
    print(f"[ingest] {estimate['page_count']} pages  "
          f"{estimate['size_mb']}MB  "
          f"profile={profile}  "
          f"~{estimate['estimated_secs']}s")

    t0    = time.time()

    # Step 1: Extract pages
    print(f"[ingest] Extracting pages...")
    pages = extract_pages(filepath)

    if not pages:
        return {"success": False, "error": "Could not extract text"}

    actual_pages = len(pages)
    full_text    = "\n\n".join(text for _, text in pages)

    if preview:
        print(f"\n── Text preview (first 500 chars) ──")
        print(full_text[:500])
        print("─" * 40 + "\n")

    # Step 2: Store document record
    stat = filepath.stat()
    now  = time.strftime("%Y-%m-%dT%H:%M:%SZ")

    conn.execute("""
        INSERT INTO documents
        (filename, filepath, file_type, file_size, page_count,
         date_modified, date_indexed, raw_text, doc_hash, status,
         created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'indexing', ?, ?)
    """, (
        filepath.name, str(filepath), filepath.suffix.lower(),
        stat.st_size, actual_pages,
        time.strftime("%Y-%m-%dT%H:%M:%SZ", time.localtime(stat.st_mtime)),
        now, full_text[:50000], doc_hash, now, now,
    ))
    conn.commit()

    doc_id = conn.execute(
        "SELECT id FROM documents WHERE doc_hash = ?", (doc_hash,)
    ).fetchone()["id"]

    if job_id:
        conn.execute(
            "UPDATE ingest_jobs SET status='processing', started_at=? WHERE id=?",
            (now, job_id)
        )
        conn.commit()

    # Step 3: Chunk
    print(f"[ingest] Chunking ({profile} profile)...")
    chunks = chunk_pages(pages, profile=profile)
    print(f"[ingest] {len(chunks)} chunks from {actual_pages} pages")

    # Step 4: Embed immediately — file is searchable from this point on
    print(f"[ingest] Embedding {len(chunks)} chunks...", end=" ", flush=True)
    try:
        n_embedded = embed_chunks(chunks, doc_id, filepath.name, conn)
        print("done")
    except Exception as e:
        print(f"ERROR: {e}")
        n_embedded = 0

    conn.execute(
        "UPDATE documents SET chunk_count=?, status='indexed' WHERE id=?",
        (len(chunks), doc_id)
    )
    conn.commit()
    print(f"[ingest] ✓ searchable — {filepath.name}")

    # Step 5: Entity extraction (enrichment pass — runs after file is already searchable)
    # Skipped for sparse/slide content and very large documents.
    total_entities = 0
    avg_chars = (len(full_text) / actual_pages) if actual_pages > 0 else 0

    if actual_pages > MAX_ENTITY_PAGES:
        print(f"[ingest] Skipping entity extraction "
              f"({actual_pages} pages > {MAX_ENTITY_PAGES} limit — large document)")
    elif avg_chars < SPARSE_THRESHOLD:
        print(f"[ingest] Skipping entity extraction "
              f"(avg {avg_chars:.0f} chars/page — slide/sparse content)")
    else:
        print(f"[ingest] Enriching: extracting entities "
              f"[{actual_pages}p, avg {avg_chars:.0f} chars/page, model={ENTITY_MODEL}]...")
        batches = [
            pages[i:i + ENTITY_BATCH_SIZE]
            for i in range(0, len(pages), ENTITY_BATCH_SIZE)
        ]

        for batch_num, batch in enumerate(batches):
            page_nums = [p for p, _ in batch]
            print(f"  batch {batch_num + 1}/{len(batches)} "
                  f"(pages {page_nums})...", end=" ", flush=True)
            n = extract_entities_from_batch(batch, doc_id, conn)
            total_entities += n
            print(f"{n} entities")

            if job_id:
                conn.execute(
                    "UPDATE ingest_jobs SET progress_page=? WHERE id=?",
                    (page_nums[-1], job_id)
                )
            conn.commit()

    # Step 6: Generate suggested questions for this document
    print(f"[ingest] Generating questions for {filepath.name}...")
    generate_doc_questions(filepath.name, chunks, conn, doc_id, model=ENTITY_MODEL)

    # Step 7: Finalise
    elapsed = round(time.time() - t0, 1)

    conn.execute(
        "UPDATE documents SET entity_count=? WHERE id=?",
        (total_entities, doc_id)
    )
    conn.commit()

    if job_id:
        conn.execute("""
            UPDATE ingest_jobs
            SET status='done', finished_at=?, progress_page=?
            WHERE id=?
        """, (time.strftime("%Y-%m-%dT%H:%M:%SZ"), actual_pages, job_id))
        conn.commit()

    result = {
        "success":  True,
        "skipped":  False,
        "doc_id":   doc_id,
        "filename": filepath.name,
        "pages":    actual_pages,
        "chunks":   len(chunks),
        "entities": total_entities,
        "embedded": n_embedded,
        "elapsed":  elapsed,
        "profile":  profile,
    }

    print(f"\n[ingest] ✓ {filepath.name}")
    print(f"  pages:    {actual_pages}")
    print(f"  chunks:   {len(chunks)}")
    print(f"  entities: {total_entities}")
    print(f"  time:     {elapsed}s")

    if preview and total_entities > 0:
        print(f"\n── Sample entities ──")
        rows = conn.execute(
            "SELECT entity_type, value, page_number FROM entities "
            "WHERE doc_id=? LIMIT 12", (doc_id,)
        ).fetchall()
        for r in rows:
            print(f"  [p.{r['page_number']}] [{r['entity_type']}] {r['value']}")

    return result


def ingest_folder(folder, conn, preview=False):
    folder = Path(folder)
    files  = []
    for ext in SUPPORTED:
        files.extend(folder.rglob(f"*{ext}"))

    if not files:
        print(f"[ingest] No supported files in {folder}")
        return []

    jobs = []
    for f in files:
        est = estimate_job(f)
        if "error" not in est:
            jobs.append((est["priority"], f, est))

    jobs.sort(key=lambda x: x[0])

    print(f"[ingest] Found {len(jobs)} files — processing smallest first\n")

    results = []
    for priority, filepath, est in jobs:
        print(f"[p{priority}] {est['filename']} "
              f"({est['page_count']} pages, ~{est['estimated_secs']}s)")
        result = ingest_file(filepath, conn, preview=preview)
        results.append(result)
        print()

    success = sum(1 for r in results if r.get("success"))
    print(f"[ingest] Done: {success}/{len(results)} files indexed")
    return results


def list_documents(conn):
    rows = conn.execute("""
        SELECT id, filename, file_type, file_size, page_count,
               chunk_count, entity_count, date_indexed, status
        FROM documents ORDER BY id DESC
    """).fetchall()
    return [dict(r) for r in rows]


# ── CLI ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    p = argparse.ArgumentParser(
        description="DONE ingest — read and index documents"
    )
    p.add_argument("--file",     help="Single file to ingest")
    p.add_argument("--folder",   help="Folder to ingest recursively")
    p.add_argument("--preview",  action="store_true",
                   help="Print extracted text and entities")
    p.add_argument("--list",     action="store_true",
                   help="List all indexed documents")
    p.add_argument("--estimate", help="Estimate ingestion time for a file")
    args = p.parse_args()

    conn = init_db()

    if args.list:
        docs = list_documents(conn)
        if not docs:
            print("[ingest] No documents indexed yet")
        else:
            print(f"\n{'ID':<4} {'File':<40} {'Pages':<7} "
                  f"{'Chunks':<8} {'Entities':<10} Status")
            print("─" * 80)
            for d in docs:
                print(f"{d['id']:<4} {d['filename']:<40} "
                      f"{d['page_count']:<7} {d['chunk_count']:<8} "
                      f"{d['entity_count']:<10} {d['status']}")

    elif args.estimate:
        est = estimate_job(args.estimate)
        if "error" in est:
            print(f"[ingest] {est['error']}")
        else:
            mins = est['estimated_secs'] // 60
            secs = est['estimated_secs'] % 60
            labels = {1: "highest", 2: "normal", 3: "low", 4: "lowest"}
            print(f"\n── Estimate: {est['filename']} ──")
            print(f"  Pages:     {est['page_count']}")
            print(f"  Size:      {est['size_mb']} MB")
            print(f"  Priority:  {est['priority']} ({labels[est['priority']]})")
            print(f"  Profile:   {est['chunk_profile']} chunks")
            print(f"  Est. time: {mins}m {secs}s")

    elif args.file:
        ingest_file(args.file, conn, preview=args.preview)

    elif args.folder:
        ingest_folder(args.folder, conn, preview=args.preview)

    else:
        p.print_help()